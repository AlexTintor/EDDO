from docx import Document
from docx2pdf import convert
import os

def generar_constancia(datos,nombreDoc):
    ruta_docx = r"C:\\VisualStudio\\Python\\EDDO\\EDDO\\EDDO\\Backend\\documentos\\{nombreDoc}.docx".format(nombreDoc=nombreDoc)
    salida_docx = r"C:\\VisualStudio\\Python\\EDDO\\EDDO\\EDDO\\Backend\\documentos\\temp_{nombreDoc}.docx".format(nombreDoc=nombreDoc)
    salida_pdf = r"C:\\VisualStudio\\Python\\EDDO\\EDDO\\EDDO\\Backend\\documentos\\{nombreDoc}.pdf".format(nombreDoc=nombreDoc)
    # Leer plantilla Word
    doc = Document(ruta_docx)

    # Función auxiliar que reemplaza texto en todos los párrafos (sin perder formato)
    def reemplazar_texto(documento, buscar, reemplazar):
        for p in documento.paragraphs:
            if buscar in p.text:
                for run in p.runs:
                    if buscar in run.text:
                        run.text = run.text.replace(buscar, reemplazar)
        for tabla in documento.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    reemplazar_texto(celda, buscar, reemplazar)

    # Reemplazar todas las variables
    for clave, valor in datos.items():
        reemplazar_texto(doc, clave, valor)

    # Guardar el nuevo documento
    doc.save(salida_docx)

    # Convertir el Word a PDF (mantiene formato e imágenes)
    convert(salida_docx, salida_pdf)
    if os.path.exists(salida_docx):
        os.remove(salida_docx)
        print("🗑️ Archivo eliminado correctamente.")
    else:
        print("⚠️ El archivo no existe:", salida_docx)

    print(f"✅ PDF generado correctamente en:\n{salida_pdf}")

    return salida_pdf


datos = {
    "VARIABLE_DOCENTE": "Juan Pérez López",
    "VARIABLE_FILACION": "123456",
    "VAR_FECHA_CONTRATACION": "15 de marzo de 2018",
    "VAR_CLAVE_PRESUNTUAL": "C12345",
    "VAR_FECHA_ESCRITA": "9 de junio de 2025",
    "VAR_DIA_MES": "1 de enero de 2024",
    "VAR_FECHA_ESC_2": "9 de junio de 2025"
}

generar_constancia(datos, "doc1")
