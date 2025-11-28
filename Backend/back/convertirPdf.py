import json
from docx import Document
from docx2pdf import convert
from docx.shared import Inches
import os

def generar_constancia(datos,nombreDoc,datos2):

    # Si datos es string y parece JSON → convertir
    if isinstance(datos, str):
        try:
            datos = json.loads(datos)
        except:
            # Convertir formato texto1:valor1; texto2:valor2
            dic = {}
            partes = datos.split(";")
            for p in partes:
                if ":" in p:
                    k, v = p.split(":", 1)
                    dic[k.strip()] = v.strip()
            datos = dic
            
    ruta_docx = r"C:\\VisualStudio\\Python\\EDDO\\EDDO\\EDDO\\documentos\\{nombreDoc}.docx".format(nombreDoc=nombreDoc)
    salida_docx = r"C:\\VisualStudio\\Python\\EDDO\\EDDO\\EDDO\\documentos\\temp_{nombreDoc}.docx".format(nombreDoc=nombreDoc)
    salida_pdf = r"C:\\VisualStudio\\Python\\EDDO\\EDDO\\EDDO\\documentos\\{nombreDoc}.pdf".format(nombreDoc=nombreDoc)
    # Leer plantilla Word
    doc = Document(ruta_docx)
    def reemplazar_imagen_en_documento(doc, marcador, ruta_imagen):
        # Buscar en párrafos
        for p in doc.paragraphs:
            if marcador in p.text:
                # Limpiar runs
                for run in p.runs:
                    run.text = ""

                # Colocar la imagen en un run NUEVO
                run = p.add_run()
                run.add_picture(ruta_imagen, width=Inches(1.5))
                print(f"🖼️ Imagen reemplazada (párrafo): {marcador}")

        # Buscar en tablas
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for p in celda.paragraphs:
                        if marcador in p.text:
                            # Limpiar runs
                            for run in p.runs:
                                run.text = ""

                            run = p.add_run()
                            run.add_picture(ruta_imagen, width=Inches(1.5))
                            print(f"🖼️ Imagen reemplazada (tabla): {marcador}")


    # Función auxiliar que reemplaza texto en todos los párrafos (sin perder formato)
    def reemplazar_en_parrafo(parrafo, buscar, reemplazar):
        if buscar in parrafo.text:
            nuevo_texto = parrafo.text.replace(buscar, reemplazar)

            # Vaciar todos los runs
            for run in parrafo.runs:
                run.text = ""

            # Escribir el nuevo texto en el primer run
            parrafo.runs[0].text = nuevo_texto


    def reemplazar_en_documento(doc, buscar, reemplazar):
        # Reemplazar en párrafos normales
        for p in doc.paragraphs:
            reemplazar_en_parrafo(p, buscar, reemplazar)

        # Reemplazar dentro de tablas
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for p in celda.paragraphs:
                        reemplazar_en_parrafo(p, buscar, reemplazar)

                    # También buscar tablas dentro de celdas
                    for subtabla in celda.tables:
                        for subfila in subtabla.rows:
                            for subcelda in subfila.cells:
                                for p2 in subcelda.paragraphs:
                                    reemplazar_en_parrafo(p2, buscar, reemplazar)

    for clave, valor in datos2.items():
        print(f"Reemplazando {clave} por {valor}")

    
    for clave, valor in datos.items():
        print(f"Reemplazando {clave} por {valor}")
        reemplazar_en_documento(doc, clave, valor)

    # Guardar el nuevo documento
    doc.save(salida_docx)

    # Convertir el Word a PDF (mantiene formato e imágenes)
    convert(salida_docx, salida_pdf)
    if os.path.exists(salida_docx):
        os.remove(salida_docx)
        print("🗑️ Archivo eliminado correctamente.")
    else:
        print("⚠️ El archivo no existe:", salida_docx)

    print(f"✅ PDF generado correctamente en:\n{salida_pdf}")

    return salida_pdf


datos = {
    "VARIABLE_DOCENTE": "Juan Pérez López",
    "VARIABLE_FILACION": "123456",
    "VAR_FECHA_CONTRATACION": "15 de marzo de 2018",
    "VAR_CLAVE_PRESUNTUAL": "C12345",
    "VAR_FECHA_ESCRITA": "9 de junio de 2025",
    "VAR_DIA_MES": "1 de enero de 2024",
    "VAR_FECHA_ESC_2": "9 de junio de 2025"
}
#generar_constancia(datos, "Constancia de Participación 1.4.8.2")